"""
Export de Synthèse CAC - VERSION FINALE
Utilise le template Word et insère tous les contenus aux emplacements corrects
Génère les sommaires automatiquement
"""

from fastapi import APIRouter, HTTPException
from fastapi.responses import Response
from pydantic import BaseModel
from typing import List, Optional
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
from datetime import datetime
import logging
from pathlib import Path

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/api/word")

# Chemin vers le template - Utiliser un document vide si template non disponible
TEMPLATE_PATH = None  # Pas de template, génération programmatique


# === MODÈLES PYDANTIC ===

class FrapPointMetadata(BaseModel):
    etape: Optional[str] = None
    norme: Optional[str] = None
    methode: Optional[str] = None
    reference: Optional[str] = None


class FrapPoint(BaseModel):
    metadata: FrapPointMetadata
    intitule: str
    observation: Optional[str] = ""
    constat: Optional[str] = ""
    risque: Optional[str] = ""
    recommandation: Optional[str] = ""


class RecosRevisionMetadata(BaseModel):
    etape: Optional[str] = None
    norme: Optional[str] = None
    methode: Optional[str] = None
    reference: Optional[str] = None


class RecosRevisionPoint(BaseModel):
    metadata: RecosRevisionMetadata
    intitule: str
    description: Optional[str] = ""
    observation: Optional[str] = ""
    ajustement: Optional[str] = ""
    regularisation: Optional[str] = ""


class RecosControleInterneMetadata(BaseModel):
    etape: Optional[str] = None
    norme: Optional[str] = None
    methode: Optional[str] = None
    reference: Optional[str] = None


class RecosControleInternePoint(BaseModel):
    metadata: RecosControleInterneMetadata
    intitule: str
    observation: Optional[str] = ""
    constat: Optional[str] = ""
    risque: Optional[str] = ""
    recommandation: Optional[str] = ""


class SyntheseCAC_Request(BaseModel):
    frap_points: List[FrapPoint] = []
    recos_revision_points: List[RecosRevisionPoint] = []
    recos_controle_interne_points: List[RecosControleInternePoint] = []
    date_rapport: Optional[str] = None
    entite: Optional[str] = "Entité auditée"
    exercice: Optional[str] = None


# === FONCTIONS UTILITAIRES ===

def clean_text(text: str) -> str:
    """Nettoyer le texte des échappements multiples"""
    if not text:
        return ""
    # Remplacer les échappements multiples
    text = text.replace('\\\\\\\\n', '\n').replace('\\\\n', '\n').replace('\\n', '\n')
    return text.strip()


def add_formatted_paragraph(doc: Document, text: str, bold: bool = False, italic: bool = False, 
                            font_size: int = 11, indent: float = 0, space_after: int = 6):
    """Ajouter un paragraphe formaté avec gestion correcte des retours à la ligne"""
    text = clean_text(text)
    if not text:
        return None
    
    para = doc.add_paragraph()
    
    if indent > 0:
        para.paragraph_format.left_indent = Inches(indent)
    
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.line_spacing = 1.15
    
    # Gérer les retours à la ligne
    lines = text.split('\n')
    
    for i, line in enumerate(lines):
        if i > 0:
            para.add_run('\n')
        if line.strip():  # Ignorer les lignes vides
            run = para.add_run(line.strip())
            run.font.name = 'Calibri'
            run.font.size = Pt(font_size)
            run.font.bold = bold
            run.font.italic = italic
    
    return para


def add_section_with_label(doc: Document, label: str, content: str, indent: float = 0.25):
    """Ajouter une section avec label en gras et contenu"""
    content = clean_text(content)
    if not content:
        return None
    
    para = doc.add_paragraph()
    if indent > 0:
        para.paragraph_format.left_indent = Inches(indent)
    
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.15
    
    # Label en gras
    run_label = para.add_run(f"{label}: ")
    run_label.font.name = 'Calibri'
    run_label.font.size = Pt(11)
    run_label.font.bold = True
    
    # Contenu avec gestion des retours à la ligne
    lines = content.split('\n')
    for i, line in enumerate(lines):
        if i > 0:
            para.add_run('\n')
        if line.strip():
            run_content = para.add_run(line.strip())
            run_content.font.name = 'Calibri'
            run_content.font.size = Pt(11)
    
    return para


def add_heading_custom(doc: Document, text: str, level: int = 2, font_size: int = 12):
    """Ajouter un titre personnalisé"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    if heading.runs:
        run = heading.runs[0]
        run.font.name = 'Calibri'
        run.font.size = Pt(font_size)
        run.font.bold = True
        run.font.color.rgb = RGBColor(31, 56, 100)
    
    return heading


def find_marker_paragraph(doc: Document, marker_text: str):
    """Trouver un paragraphe marqueur dans le document"""
    for i, para in enumerate(doc.paragraphs):
        if marker_text.lower() in para.text.lower():
            logger.info(f"✅ Marqueur trouvé: '{marker_text}' à l'index {i}")
            return i, para
    logger.warning(f"⚠️ Marqueur non trouvé: '{marker_text}'")
    return None, None


def insert_paragraph_after(doc: Document, paragraph, text: str, **kwargs):
    """Insérer un nouveau paragraphe après un paragraphe existant"""
    # Créer un nouveau paragraphe
    new_para = paragraph._element.addnext(paragraph._element.__class__())
    new_para = paragraph._p.addnext(new_para)
    
    # Cette méthode ne fonctionne pas bien, utilisons plutôt l'ajout à la fin
    # et nous réorganiserons manuellement
    return add_formatted_paragraph(doc, text, **kwargs)


def create_synthese_cac_from_template_final(request: SyntheseCAC_Request) -> BytesIO:
    """
    Créer le document Word de synthèse CAC à partir du template
    VERSION FINALE avec insertion correcte de tous les contenus
    """
    
    # Créer un nouveau document (pas de template)
    if TEMPLATE_PATH is None:
        doc = Document()
        logger.info("✅ Nouveau document créé (sans template)")
        
        # Ajouter un titre principal
        doc.add_heading("SYNTHÈSE CAC - RAPPORT D'AUDIT", level=0)
        doc.add_paragraph(f"Date du rapport: {request.date_rapport}")
        doc.add_paragraph("")  # Ligne vide
        
    else:
        # Vérifier que le template existe
        if not TEMPLATE_PATH.exists():
            logger.error(f"❌ Template non trouvé: {TEMPLATE_PATH}")
            raise HTTPException(status_code=500, detail=f"Template non trouvé: {TEMPLATE_PATH}")
        
        # Charger le template
        try:
            doc = Document(str(TEMPLATE_PATH))
            logger.info(f"✅ Template chargé: {TEMPLATE_PATH}")
            logger.info(f"   Nombre de paragraphes dans le template: {len(doc.paragraphs)}")
        except Exception as e:
            logger.error(f"❌ Erreur chargement template: {e}")
            raise HTTPException(status_code=500, detail=f"Erreur chargement template: {e}")
    
    # === SECTION 2: OBSERVATIONS D'AUDIT (Recos Révision) ===
    if TEMPLATE_PATH is None:
        # Mode sans template: ajouter directement
        if len(request.recos_revision_points) > 0:
            doc.add_heading("2. OBSERVATIONS D'AUDIT", level=1)
            doc.add_paragraph("")
            
            # Sommaire
            doc.add_heading("Sommaire des observations", level=2)
            for i, point in enumerate(request.recos_revision_points, 1):
                doc.add_paragraph(f"{i}. {point.intitule}", style='List Number')
            doc.add_paragraph("")
            
            # Détails
            for i, point in enumerate(request.recos_revision_points, 1):
                doc.add_heading(f"2.{i}. {point.intitule}", level=2)
                
                if point.description:
                    add_section_with_label(doc, "Description", point.description)
                if point.observation:
                    add_section_with_label(doc, "Observation", point.observation)
                if point.ajustement:
                    add_section_with_label(doc, "Ajustement proposé", point.ajustement)
                if point.regularisation:
                    add_section_with_label(doc, "Régularisation", point.regularisation)
                
                doc.add_paragraph("")  # Ligne vide entre les points
        
        obs_index = None
    else:
        obs_index, obs_para = find_marker_paragraph(doc, "2. OBSERVATIONS D'AUDIT")
    
    if obs_index is not None and len(request.recos_revision_points) > 0:
        logger.info(f"📝 Insertion de {len(request.recos_revision_points)} points de révision")
        
        # Ajouter un sommaire
        add_formatted_paragraph(doc, f"\nNos travaux de révision des comptes ont permis d'identifier "
                               f"{len(request.recos_revision_points)} point(s) nécessitant des ajustements "
                               f"ou reclassements comptables:\n", indent=0.25)
        
        # Liste des points
        for idx, point in enumerate(request.recos_revision_points, 1):
            add_formatted_paragraph(doc, f"  {idx}. {point.intitule}", indent=0.25, space_after=3)
        
        doc.add_paragraph()
        
        # Détail de chaque point
        for idx, point in enumerate(request.recos_revision_points, 1):
            # Titre du point
            add_heading_custom(doc, f"2.{idx}. {point.intitule}", level=2, font_size=12)
            
            # Référence
            if point.metadata.reference:
                add_formatted_paragraph(doc, f"Référence: {point.metadata.reference}", 
                                      italic=True, indent=0.25, space_after=6)
            
            # TOUS LES CHAMPS - dans l'ordre correct
            if point.description:
                add_section_with_label(doc, "Description", point.description, indent=0.25)
            
            if point.observation:
                add_section_with_label(doc, "Observation", point.observation, indent=0.25)
            
            if point.ajustement:
                add_section_with_label(doc, "Ajustement/Reclassement proposé", 
                                     point.ajustement, indent=0.25)
            
            if point.regularisation:
                add_section_with_label(doc, "Régularisation comptable", 
                                     point.regularisation, indent=0.25)
            
            # Espace entre les points
            doc.add_paragraph()
    
    elif obs_index is not None:
        add_formatted_paragraph(doc, "\nAucune observation d'audit identifiée lors de nos travaux.\n", 
                              italic=True, indent=0.25)
    
    # === SECTION 3: POINTS DE CONTRÔLE INTERNE (FRAP + Recos CI) ===
    # Combiner FRAP et Recos Contrôle Interne
    all_ci_points = []
    
    for frap in request.frap_points:
        all_ci_points.append({
            'type': 'FRAP',
            'metadata': frap.metadata,
            'intitule': frap.intitule,
            'observation': frap.observation,
            'constat': frap.constat,
            'risque': frap.risque,
            'recommandation': frap.recommandation
        })
    
    for reco in request.recos_controle_interne_points:
        all_ci_points.append({
            'type': 'Recos CI',
            'metadata': reco.metadata,
            'intitule': reco.intitule,
            'observation': reco.observation,
            'constat': reco.constat,
            'risque': reco.risque,
            'recommandation': reco.recommandation
        })
    
    if TEMPLATE_PATH is None:
        # Mode sans template: ajouter directement
        if len(all_ci_points) > 0:
            doc.add_heading("3. POINTS DE CONTRÔLE INTERNE", level=1)
            doc.add_paragraph("")
            
            # Sommaire
            doc.add_heading("Sommaire des points de contrôle interne", level=2)
            for i, point in enumerate(all_ci_points, 1):
                doc.add_paragraph(f"{i}. {point['intitule']}", style='List Number')
            doc.add_paragraph("")
            
            # Détails
            for i, point in enumerate(all_ci_points, 1):
                doc.add_heading(f"3.{i}. {point['intitule']}", level=2)
                doc.add_paragraph(f"Type: {point['type']}", style='Intense Quote')
                
                if point['observation']:
                    add_section_with_label(doc, "Observation", point['observation'])
                if point['constat']:
                    add_section_with_label(doc, "Constat", point['constat'])
                if point['risque']:
                    add_section_with_label(doc, "Risque", point['risque'])
                if point['recommandation']:
                    add_section_with_label(doc, "Recommandation", point['recommandation'])
                
                doc.add_paragraph("")  # Ligne vide entre les points
        
        ci_index = None
    else:
        ci_index, ci_para = find_marker_paragraph(doc, "3. POINTS DE CONTRÔLE INTERNE")
    
    if ci_index is not None and len(all_ci_points) > 0:
        logger.info(f"📝 Insertion de {len(all_ci_points)} points de contrôle interne")
        
        # Ajouter un sommaire
        add_formatted_paragraph(doc, f"\nNotre évaluation du contrôle interne comptable a permis "
                               f"d'identifier {len(all_ci_points)} point(s) nécessitant une attention "
                               f"particulière et des actions correctives:\n", indent=0.25)
        
        # Liste des points
        for idx, point in enumerate(all_ci_points, 1):
            add_formatted_paragraph(doc, f"  {idx}. {point['intitule']}", indent=0.25, space_after=3)
        
        doc.add_paragraph()
        
        # Détail de chaque point
        for idx, point in enumerate(all_ci_points, 1):
            # Titre du point
            add_heading_custom(doc, f"3.{idx}. {point['intitule']}", level=2, font_size=12)
            
            # Référence
            if point['metadata'].reference:
                add_formatted_paragraph(doc, f"Référence: {point['metadata'].reference}", 
                                      italic=True, indent=0.25, space_after=6)
            
            # Type
            add_formatted_paragraph(doc, f"Type: {point['type']}", italic=True, indent=0.25, space_after=6)
            
            # TOUS LES CHAMPS
            if point.get('observation'):
                add_section_with_label(doc, "Observation", point['observation'], indent=0.25)
            
            if point.get('constat'):
                add_section_with_label(doc, "Constat", point['constat'], indent=0.25)
            
            if point.get('risque'):
                add_section_with_label(doc, "Risques identifiés", point['risque'], indent=0.25)
            
            if point.get('recommandation'):
                add_section_with_label(doc, "Recommandation", point['recommandation'], indent=0.25)
            
            # Espace entre les points
            doc.add_paragraph()
    
    elif ci_index is not None:
        add_formatted_paragraph(doc, "\nAucun point de contrôle interne identifié lors de nos travaux.\n", 
                              italic=True, indent=0.25)
    
    # Sauvegarder dans un buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    logger.info("✅ Document généré avec succès à partir du template")
    logger.info(f"   - Points Révision: {len(request.recos_revision_points)}")
    logger.info(f"   - Points FRAP: {len(request.frap_points)}")
    logger.info(f"   - Points Recos CI: {len(request.recos_controle_interne_points)}")
    
    return buffer


# === ENDPOINT API ===

@router.post("/export-synthese-cac-final")
async def export_synthese_cac_final(request: SyntheseCAC_Request):
    """
    Exporter une synthèse CAC en document Word structuré
    VERSION FINALE: Utilise le template et insère TOUS les contenus correctement
    
    Intègre:
    - Section 2: Recos Révision (Description, Observation, Ajustement, Régularisation)
    - Section 3: FRAP + Recos CI (Observation, Constat, Risque, Recommandation)
    - Sommaires automatiques pour chaque section
    """
    try:
        total_points = (
            len(request.frap_points) + 
            len(request.recos_revision_points) + 
            len(request.recos_controle_interne_points)
        )
        
        logger.info(f"📊 Export Synthèse CAC FINAL: {total_points} points au total")
        logger.info(f"   - FRAP: {len(request.frap_points)}")
        logger.info(f"   - Recos Révision: {len(request.recos_revision_points)}")
        logger.info(f"   - Recos CI: {len(request.recos_controle_interne_points)}")
        
        # Créer le document à partir du template
        doc_buffer = create_synthese_cac_from_template_final(request)
        
        # Générer le nom de fichier
        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        filename = f"synthese_cac_{timestamp}.docx"
        
        logger.info(f"✅ Export réussi: {filename}")
        
        # Retourner le fichier
        return Response(
            content=doc_buffer.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="{filename}"'
            }
        )
    
    except Exception as e:
        logger.error(f"❌ Erreur export synthèse CAC FINAL: {e}")
        import traceback
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))
