"""
Export de Synthèse CAC - Version 2 Simplifiée
Génère des rapports structurés SANS utiliser de template
Crée le document programmatiquement avec la structure exacte du template
"""

from fastapi import APIRouter, HTTPException
from fastapi.responses import Response
from pydantic import BaseModel
from typing import List, Optional
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from datetime import datetime
import logging

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/api/word")


# === MODÈLES PYDANTIC ===

class FrapPointMetadata(BaseModel):
    etape: Optional[str] = None
    norme: Optional[str] = None
    methode: Optional[str] = None
    reference: Optional[str] = None


class FrapPoint(BaseModel):
    metadata: FrapPointMetadata
    intitule: str
    observation: Optional[str] = ""
    constat: Optional[str] = ""
    risque: Optional[str] = ""
    recommandation: Optional[str] = ""


class RecosRevisionMetadata(BaseModel):
    etape: Optional[str] = None
    norme: Optional[str] = None
    methode: Optional[str] = None
    reference: Optional[str] = None


class RecosRevisionPoint(BaseModel):
    metadata: RecosRevisionMetadata
    intitule: str
    description: Optional[str] = ""
    observation: Optional[str] = ""
    ajustement: Optional[str] = ""
    regularisation: Optional[str] = ""


class RecosControleInterneMetadata(BaseModel):
    etape: Optional[str] = None
    norme: Optional[str] = None
    methode: Optional[str] = None
    reference: Optional[str] = None


class RecosControleInternePoint(BaseModel):
    metadata: RecosControleInterneMetadata
    intitule: str
    observation: Optional[str] = ""
    constat: Optional[str] = ""
    risque: Optional[str] = ""
    recommandation: Optional[str] = ""


class SyntheseCAC_Request(BaseModel):
    frap_points: List[FrapPoint] = []
    recos_revision_points: List[RecosRevisionPoint] = []
    recos_controle_interne_points: List[RecosControleInternePoint] = []
    date_rapport: Optional[str] = None
    entite: Optional[str] = "Entité auditée"
    exercice: Optional[str] = None


# === FONCTIONS UTILITAIRES ===

def add_paragraph_with_style(doc, text, bold=False, italic=False, font_size=11, 
                             indent=0, space_after=6, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """Ajouter un paragraphe avec style personnalisé"""
    para = doc.add_paragraph()
    para.alignment = alignment
    
    if indent > 0:
        para.paragraph_format.left_indent = Inches(indent)
    
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.line_spacing = 1.15
    
    # Gérer les retours à la ligne
    lines = text.replace('\\\\n', '\n').replace('\\n', '\n').split('\n')
    
    for i, line in enumerate(lines):
        if i > 0:
            para.add_run('\n')
        run = para.add_run(line.strip())
        run.font.name = 'Calibri'
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.italic = italic
    
    return para


def add_labeled_content(doc, label, content, indent=0.25):
    """Ajouter un contenu avec label en gras"""
    if not content or content.strip() == "":
        return
    
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(indent)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.15
    
    # Label en gras
    run_label = para.add_run(f"{label}: ")
    run_label.font.name = 'Calibri'
    run_label.font.size = Pt(11)
    run_label.font.bold = True
    
    # Contenu avec retours à la ligne
    lines = content.replace('\\\\n', '\n').replace('\\n', '\n').split('\n')
    for i, line in enumerate(lines):
        if i > 0:
            para.add_run('\n')
        run_content = para.add_run(line.strip())
        run_content.font.name = 'Calibri'
        run_content.font.size = Pt(11)


def add_custom_heading(doc, text, level=1, font_size=14, color_rgb=(31, 56, 100)):
    """Ajouter un titre personnalisé"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    run = heading.runs[0]
    run.font.name = 'Calibri'
    run.font.size = Pt(font_size)
    run.font.bold = True
    run.font.color.rgb = RGBColor(*color_rgb)
    
    return heading


def create_synthese_cac_document_v2(request: SyntheseCAC_Request) -> BytesIO:
    """
    Créer le document Word de synthèse CAC programmatiquement
    Structure identique au template
    """
    
    doc = Document()
    
    # Configuration des marges
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # === EN-TÊTE ===
    add_custom_heading(doc, 'SYNTHÈSE DES TRAVAUX DE RÉVISION', level=0, font_size=16)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    if request.entite:
        add_paragraph_with_style(doc, f"Entité: {request.entite}", bold=True)
    if request.exercice:
        add_paragraph_with_style(doc, f"Exercice: {request.exercice}", bold=True)
    if request.date_rapport:
        add_paragraph_with_style(doc, f"Date du rapport: {request.date_rapport}", bold=True)
    
    doc.add_paragraph()
    
    # === 1. INTRODUCTION ===
    add_custom_heading(doc, "1. INTRODUCTION", level=1, font_size=14)
    
    intro_text = (
        "Le présent document synthétise les observations et recommandations issues de nos travaux "
        "de révision des comptes et d'évaluation du contrôle interne comptable. "
        "Ces travaux ont été réalisés conformément aux normes professionnelles applicables "
        "et visent à identifier les ajustements comptables nécessaires ainsi que les améliorations "
        "à apporter au dispositif de contrôle interne."
    )
    add_paragraph_with_style(doc, intro_text)
    
    doc.add_paragraph()
    
    # === 2. OBSERVATIONS D'AUDIT ===
    add_custom_heading(doc, "2. OBSERVATIONS D'AUDIT", level=1, font_size=14)
    
    if len(request.recos_revision_points) == 0:
        add_paragraph_with_style(doc, "Aucune observation d'audit identifiée.", italic=True, indent=0.25)
    else:
        intro_obs = (
            f"Nos travaux de révision des comptes ont permis d'identifier {len(request.recos_revision_points)} "
            f"point(s) nécessitant des ajustements ou reclassements comptables."
        )
        add_paragraph_with_style(doc, intro_obs)
        doc.add_paragraph()
        
        for idx, point in enumerate(request.recos_revision_points, 1):
            # Titre du point
            add_custom_heading(doc, f"2.{idx}. {point.intitule}", level=2, font_size=12)
            
            # Référence
            if point.metadata.reference:
                add_paragraph_with_style(doc, f"Référence: {point.metadata.reference}", 
                                       italic=True, indent=0.25)
            
            # TOUS LES CHAMPS
            if point.description:
                add_labeled_content(doc, "Description", point.description, indent=0.25)
            
            if point.observation:
                add_labeled_content(doc, "Observation", point.observation, indent=0.25)
            
            if point.ajustement:
                add_labeled_content(doc, "Ajustement/Reclassement proposé", 
                                  point.ajustement, indent=0.25)
            
            if point.regularisation:
                add_labeled_content(doc, "Régularisation comptable", 
                                  point.regularisation, indent=0.25)
            
            doc.add_paragraph()
    
    doc.add_paragraph()
    
    # === 3. POINTS DE CONTRÔLE INTERNE ===
    add_custom_heading(doc, "3. POINTS DE CONTRÔLE INTERNE", level=1, font_size=14)
    
    # Combiner FRAP et Recos CI
    all_ci_points = []
    
    for frap in request.frap_points:
        all_ci_points.append({
            'type': 'FRAP',
            'metadata': frap.metadata,
            'intitule': frap.intitule,
            'observation': frap.observation,
            'constat': frap.constat,
            'risque': frap.risque,
            'recommandation': frap.recommandation
        })
    
    for reco in request.recos_controle_interne_points:
        all_ci_points.append({
            'type': 'Recos CI',
            'metadata': reco.metadata,
            'intitule': reco.intitule,
            'observation': reco.observation,
            'constat': reco.constat,
            'risque': reco.risque,
            'recommandation': reco.recommandation
        })
    
    if len(all_ci_points) == 0:
        add_paragraph_with_style(doc, "Aucun point de contrôle interne identifié.", 
                               italic=True, indent=0.25)
    else:
        intro_ci = (
            f"Notre évaluation du contrôle interne comptable a permis d'identifier {len(all_ci_points)} "
            f"point(s) nécessitant une attention particulière et des actions correctives."
        )
        add_paragraph_with_style(doc, intro_ci)
        doc.add_paragraph()
        
        for idx, point in enumerate(all_ci_points, 1):
            # Titre du point
            add_custom_heading(doc, f"3.{idx}. {point['intitule']}", level=2, font_size=12)
            
            # Référence
            if point['metadata'].reference:
                add_paragraph_with_style(doc, f"Référence: {point['metadata'].reference}", 
                                       italic=True, indent=0.25)
            
            # Type
            add_paragraph_with_style(doc, f"Type: {point['type']}", italic=True, indent=0.25)
            
            # Tous les champs
            if point.get('observation'):
                add_labeled_content(doc, "Observation", point['observation'], indent=0.25)
            
            if point.get('constat'):
                add_labeled_content(doc, "Constat", point['constat'], indent=0.25)
            
            if point.get('risque'):
                add_labeled_content(doc, "Risques identifiés", point['risque'], indent=0.25)
            
            if point.get('recommandation'):
                add_labeled_content(doc, "Recommandation", point['recommandation'], indent=0.25)
            
            doc.add_paragraph()
    
    # === 4. CONCLUSION ===
    doc.add_paragraph()
    add_custom_heading(doc, "4. CONCLUSION", level=1, font_size=14)
    
    total_points = len(request.recos_revision_points) + len(all_ci_points)
    conclusion_text = (
        f"Au total, {total_points} point(s) ont été identifié(s) lors de nos travaux. "
        "Nous recommandons à la Direction de mettre en œuvre les ajustements comptables "
        "et les actions correctives proposées dans les meilleurs délais. "
        "Nous restons à votre disposition pour tout complément d'information."
    )
    add_paragraph_with_style(doc, conclusion_text)
    
    # Sauvegarder
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    logger.info("✅ Document généré avec succès (version programmatique)")
    
    return buffer


# === ENDPOINT API ===

@router.post("/export-synthese-cac-v2")
async def export_synthese_cac_v2(request: SyntheseCAC_Request):
    """
    Exporter une synthèse CAC en document Word structuré
    VERSION 2: Création programmatique (sans template)
    
    Intègre TOUS les champs:
    - Recos Révision: Description, Observation, Ajustement, Régularisation
    - FRAP: Observation, Constat, Risque, Recommandation
    - Recos CI: Observation, Constat, Risque, Recommandation
    """
    try:
        total_points = (
            len(request.frap_points) + 
            len(request.recos_revision_points) + 
            len(request.recos_controle_interne_points)
        )
        
        logger.info(f"📊 Export Synthèse CAC V2: {total_points} points au total")
        logger.info(f"   - FRAP: {len(request.frap_points)}")
        logger.info(f"   - Recos Révision: {len(request.recos_revision_points)}")
        logger.info(f"   - Recos CI: {len(request.recos_controle_interne_points)}")
        
        # Créer le document
        doc_buffer = create_synthese_cac_document_v2(request)
        
        # Générer le nom de fichier
        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        filename = f"synthese_cac_{timestamp}.docx"
        
        logger.info(f"✅ Export réussi: {filename}")
        
        # Retourner le fichier
        return Response(
            content=doc_buffer.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="{filename}"'
            }
        )
    
    except Exception as e:
        logger.error(f"❌ Erreur export synthèse CAC V2: {e}")
        raise HTTPException(status_code=500, detail=str(e))
